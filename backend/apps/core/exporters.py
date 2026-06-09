"""
Generic tabular exporters — turn a title + headers + rows into a
downloadable Excel (.xlsx), Word (.docx), or PDF (.pdf) file.

Used across modules (inventory, vendors, bookings) so every list can
offer the same "download / share / save" options.
"""
import io

from django.http import HttpResponse

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

from docx import Document
from docx.shared import Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle


EXCEL_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
WORD_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
PDF_CONTENT_TYPE = 'application/pdf'

_DARK = colors.HexColor('#1a1a2e')
_LIGHT = colors.HexColor('#f8f9fa')


def _stringify(rows):
    return [[('' if cell is None else str(cell)) for cell in row] for row in rows]


def export_excel(title, headers, rows):
    """Return a BytesIO of an .xlsx with a styled header row."""
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31] or 'Sheet1'  # Excel caps sheet names at 31 chars

    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor='1A1A2E')

    ws.append(list(headers))
    for col_idx, _ in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='left')

    for row in _stringify(rows):
        ws.append(row)

    # Auto-ish column widths
    for col_idx, header in enumerate(headers, start=1):
        longest = len(str(header))
        for row in rows:
            longest = max(longest, len(str(row[col_idx - 1])) if row[col_idx - 1] is not None else 0)
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(longest + 4, 50)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def export_word(title, headers, rows):
    """Return a BytesIO of a .docx with a heading and a styled table."""
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row in _stringify(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_pdf(title, headers, rows):
    """Return a BytesIO of a landscape A4 PDF with a styled table."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    elems = [Paragraph(title, styles['Title']), Spacer(1, 0.4 * cm)]

    data = [list(headers)] + _stringify(rows)
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), _DARK),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, _LIGHT]),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elems.append(table)
    doc.build(elems)
    buf.seek(0)
    return buf


def build_export_response(fmt, filename_base, title, headers, rows):
    """
    Dispatch on fmt ('excel' | 'word' | 'pdf', default excel) and return an
    HttpResponse with the right content type and attachment filename.
    """
    fmt = (fmt or 'excel').lower()
    if fmt == 'pdf':
        buf, content_type, ext = export_pdf(title, headers, rows), PDF_CONTENT_TYPE, 'pdf'
    elif fmt in ('word', 'docx', 'doc'):
        buf, content_type, ext = export_word(title, headers, rows), WORD_CONTENT_TYPE, 'docx'
    else:
        buf, content_type, ext = export_excel(title, headers, rows), EXCEL_CONTENT_TYPE, 'xlsx'

    resp = HttpResponse(buf.read(), content_type=content_type)
    resp['Content-Disposition'] = f'attachment; filename="{filename_base}.{ext}"'
    return resp
